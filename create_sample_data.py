"""
Enterprise Knowledge Assistant - Sample Data Generator

Creates sample PDF, TXT, and DOCX files for demonstration purposes.
Run this script before ingestion to populate the sample_data/ directory.
"""

import sys
from pathlib import Path

SAMPLE_DIR = Path(__file__).parent / "sample_data"
SAMPLE_DIR.mkdir(exist_ok=True)


def create_hr_policy_pdf():
    """Create a multi-page HR Policy PDF."""
    try:
        from fpdf import FPDF
    except ImportError:
        print("ERROR: fpdf2 is not installed. Run: pip install fpdf2")
        sys.exit(1)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # ── Page 1: Overview ──────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 15, "HR Policy Document", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 8, "Acme Corporation - Effective January 2024", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "1. Company Overview", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Acme Corporation is a leading technology company specializing in enterprise software solutions. "
        "Founded in 2005, we employ over 2,000 people across 15 offices worldwide. Our mission is to "
        "empower businesses with innovative, reliable, and scalable technology. This HR Policy document "
        "outlines the key policies and benefits available to all full-time and part-time employees."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "2. Employment Classification", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Employees are classified as follows:\n"
        "- Full-Time: 40+ hours per week, eligible for all benefits.\n"
        "- Part-Time: 20-39 hours per week, eligible for prorated benefits.\n"
        "- Contract: Project-based engagement, not eligible for company benefits.\n"
        "- Intern: Temporary educational placement, stipend provided."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "3. Work Hours and Attendance", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Standard working hours are 9:00 AM to 6:00 PM, Monday through Friday. Employees are expected "
        "to maintain a minimum of 8 working hours per day. Flexible working arrangements are available "
        "with manager approval. Remote work is permitted up to 3 days per week for eligible roles."
    ))

    # ── Page 2: Leave Policy ──────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "4. Leave Policy", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Employees are eligible for 24 paid leaves annually, distributed as follows:\n\n"
        "- Casual Leave: 12 days per year\n"
        "- Sick Leave: 8 days per year\n"
        "- Personal Leave: 4 days per year\n\n"
        "Unused casual leaves can be carried forward to the next year, up to a maximum of 5 days. "
        "Sick leaves require a medical certificate if taken for 3 or more consecutive days. "
        "Employees must apply for planned leave at least 3 business days in advance through the "
        "HR portal. Emergency leave requests should be communicated to the immediate manager "
        "within 2 hours of the start of the workday."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "4.1 Maternity and Paternity Leave", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Maternity leave: 26 weeks of paid leave for the birth or adoption of a child. "
        "Paternity leave: 2 weeks of paid leave within 6 months of the child's birth. "
        "Adoption leave: 12 weeks for the primary caregiver, 2 weeks for the secondary caregiver."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "4.2 Public Holidays", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "The company observes 10 public holidays per year, as announced at the beginning of each "
        "calendar year. Employees required to work on a public holiday will receive compensatory "
        "time off or overtime pay at 1.5x their regular rate."
    ))

    # ── Page 3: Compensation & Benefits ───────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "5. Compensation and Benefits", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Salary reviews are conducted annually in March. Performance bonuses are paid quarterly "
        "based on individual and team KPIs. The company provides the following benefits:\n\n"
        "- Health Insurance: Comprehensive coverage for employee + family (spouse and 2 children). "
        "Premium is fully covered by the company.\n"
        "- Dental and Vision: Optional coverage with employee co-pay of 20%.\n"
        "- Life Insurance: 3x annual salary coverage.\n"
        "- Retirement Plan: Company matches 401(k) contributions up to 6% of salary.\n"
        "- Employee Stock Purchase Plan (ESPP): 15% discount on company stock.\n"
        "- Learning Allowance: $2,000 per year for courses, certifications, or conferences.\n"
        "- Wellness Benefit: $500 per year for gym, mental health apps, or wellness programs."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "6. Complaints and Grievances", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Employees may file a complaint through the HR portal or by emailing hr-complaints@acme.com. "
        "All complaints are reviewed within 5 business days. The company follows a strict non-retaliation "
        "policy. Complaints related to harassment or discrimination are escalated to the Ethics Committee "
        "and must be resolved within 30 calendar days. Anonymous reporting is available through the "
        "company's ethics hotline at 1-800-ACME-ETH."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "7. Data Privacy and Security", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "All employees must complete annual data privacy training. Personal data must be handled in "
        "compliance with GDPR and local data protection laws. Sharing sensitive company data externally "
        "without authorization is grounds for immediate termination. Employees must use company-approved "
        "tools for data storage and communication."
    ))

    output_path = SAMPLE_DIR / "hr_policy.pdf"
    pdf.output(str(output_path))
    print(f"Created: {output_path}")


def create_product_faq_txt():
    """Create a product FAQ text file."""
    content = """PRODUCT FAQ - Acme Enterprise Suite
=====================================

Q: What is Acme Enterprise Suite?
A: Acme Enterprise Suite is a comprehensive business management platform that includes CRM, project management, HR management, and analytics modules. It is designed for mid-to-large enterprises and supports up to 10,000 concurrent users.

Q: What are the system requirements?
A: The minimum system requirements are:
- Operating System: Windows 10/11, macOS 12+, or Ubuntu 20.04+
- RAM: 8 GB (16 GB recommended)
- Storage: 500 MB for the application
- Browser: Chrome 90+, Firefox 85+, Safari 14+, Edge 90+
- Internet: 10 Mbps minimum for cloud features

Q: How much does the software cost?
A: Pricing is as follows:
- Starter Plan: $29/user/month (up to 50 users)
- Business Plan: $59/user/month (up to 500 users)
- Enterprise Plan: Custom pricing (500+ users)
All plans include 24/7 support and a 30-day free trial.

Q: What is the refund policy?
A: Customers may request a full refund within 30 days of purchase if they are not satisfied with the product. After 30 days, refunds are prorated based on the remaining subscription period. Annual subscriptions receive a 20% discount but are non-refundable after 60 days. To request a refund, contact support@acme.com or call 1-800-ACME-SUP.

Q: How do I reset my password?
A: To reset your password:
1. Go to https://app.acme.com/reset
2. Enter your registered email address.
3. Click "Send Reset Link".
4. Check your email and click the reset link (valid for 24 hours).
5. Choose a new password (minimum 12 characters, must include uppercase, lowercase, number, and special character).

Q: Is there an API available?
A: Yes, Acme Enterprise Suite provides a RESTful API for integration. API documentation is available at https://docs.acme.com/api. API access requires a Business or Enterprise plan. Rate limits: 1000 requests/minute for Business, 5000 requests/minute for Enterprise.

Q: How do I contact support?
A: Support channels:
- Email: support@acme.com (response within 4 hours)
- Phone: 1-800-ACME-SUP (24/7 for Enterprise, business hours for other plans)
- Live Chat: Available on the app dashboard
- Knowledge Base: https://help.acme.com
- Community Forum: https://community.acme.com

Q: What integrations are available?
A: Acme Enterprise Suite integrates with:
- Slack, Microsoft Teams (communication)
- Salesforce, HubSpot (CRM)
- Jira, Asana (project management)
- QuickBooks, Xero (accounting)
- Google Workspace, Microsoft 365 (productivity)
- Zapier (custom workflows)
- Custom integrations via REST API and webhooks

Q: What security certifications does Acme have?
A: Acme Enterprise Suite is certified under:
- SOC 2 Type II
- ISO 27001
- GDPR compliant
- HIPAA compliant (Enterprise plan)
- PCI DSS Level 1
All data is encrypted at rest (AES-256) and in transit (TLS 1.3).

Q: How do I upgrade my plan?
A: To upgrade your plan:
1. Log into the admin dashboard at https://app.acme.com/admin.
2. Navigate to Billing > Subscription.
3. Select your desired plan.
4. Confirm the upgrade.
Upgrades are prorated and take effect immediately. Downgrades take effect at the end of the current billing cycle.
"""
    output_path = SAMPLE_DIR / "product_faq.txt"
    output_path.write_text(content, encoding="utf-8")
    print(f"Created: {output_path}")


def create_onboarding_docx():
    """Create an employee onboarding DOCX file."""
    try:
        from docx import Document
    except ImportError:
        print("WARNING: python-docx is not installed. Skipping DOCX creation.")
        print("Install with: pip install python-docx")
        return

    doc = Document()

    doc.add_heading("Employee Onboarding Guide", level=0)
    doc.add_paragraph(
        "Welcome to Acme Corporation! This guide will help you get started "
        "during your first weeks at the company."
    )

    doc.add_heading("1. Pre-Arrival Checklist", level=1)
    doc.add_paragraph(
        "Before your first day, please ensure the following are completed:"
    )
    items = [
        "Sign and return your offer letter and employment agreement.",
        "Complete the background verification consent form.",
        "Submit your bank details for payroll setup.",
        "Provide emergency contact information.",
        "Review the Employee Handbook (available on the HR portal).",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. First Day", level=1)
    doc.add_paragraph(
        "On your first day, report to the reception at 9:00 AM. Your manager or an HR "
        "representative will greet you and guide you through the following:"
    )
    items = [
        "Office tour and introduction to your team.",
        "IT setup: laptop, email account, VPN access, and software installations.",
        "Badge and building access provisioning.",
        "Mandatory orientation session (10:00 AM - 12:00 PM).",
        "Lunch with your team.",
        "Review of your 30-60-90 day plan with your manager.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. First Week Training", level=1)
    doc.add_paragraph(
        "During your first week, you are expected to complete the following training modules:"
    )
    items = [
        "Company Culture and Values (1 hour) - Learn about Acme's mission, vision, and core values.",
        "Data Security and Privacy Training (2 hours) - Mandatory for all employees.",
        "HR Policies Overview (1 hour) - Leave policies, benefits enrollment, code of conduct.",
        "Tools and Systems Training (3 hours) - Jira, Confluence, Slack, email, VPN.",
        "Product Overview (2 hours) - Introduction to Acme Enterprise Suite and key products.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Probation Period", level=1)
    doc.add_paragraph(
        "All new employees undergo a 90-day probation period. During this time, your manager "
        "will conduct bi-weekly check-ins to assess your progress. At the end of the probation "
        "period, a formal review will determine your confirmation. Employees who do not meet "
        "expectations may have their probation extended by 30 days or their employment terminated."
    )

    doc.add_heading("5. Key Contacts", level=1)
    doc.add_paragraph("HR Department: hr@acme.com | Ext: 1001")
    doc.add_paragraph("IT Helpdesk: it-support@acme.com | Ext: 1002")
    doc.add_paragraph("Facilities: facilities@acme.com | Ext: 1003")
    doc.add_paragraph("Your Manager: Refer to your offer letter for details.")

    output_path = SAMPLE_DIR / "onboarding.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_security_guidelines_pdf():
    """Create a data security guidelines PDF."""
    try:
        from fpdf import FPDF
    except ImportError:
        print("ERROR: fpdf2 is not installed. Run: pip install fpdf2")
        return

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # ── Page 1 ────────────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 15, "Data Security Guidelines", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 8, "Acme Corporation - Information Security Team", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "1. Password Policy", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "All employees must adhere to the following password requirements:\n\n"
        "- Minimum length: 12 characters\n"
        "- Must include: uppercase, lowercase, numbers, and special characters\n"
        "- Passwords must be changed every 90 days\n"
        "- Cannot reuse the last 10 passwords\n"
        "- Multi-factor authentication (MFA) is mandatory for all systems\n"
        "- Passwords must never be shared, written down, or stored in plain text"
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "2. Data Classification", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "All company data is classified into four categories:\n\n"
        "- Public: Marketing materials, published content. No restrictions.\n"
        "- Internal: Company communications, project plans. Accessible to all employees.\n"
        "- Confidential: Customer data, financial reports, contracts. Restricted access.\n"
        "- Restricted: Trade secrets, M&A plans, security credentials. Need-to-know basis only.\n\n"
        "Mishandling confidential or restricted data may result in disciplinary action up to "
        "and including termination."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "3. Acceptable Use Policy", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Company devices and networks must be used primarily for business purposes. "
        "Prohibited activities include:\n\n"
        "- Downloading unauthorized software\n"
        "- Accessing inappropriate or illegal content\n"
        "- Using personal cloud storage for company data\n"
        "- Connecting unauthorized devices to the company network\n"
        "- Disabling security software (antivirus, firewall, endpoint protection)"
    ))

    # ── Page 2 ────────────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "4. Incident Response", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "In the event of a data breach or security incident:\n\n"
        "1. Immediately report the incident to security@acme.com or call the Security Operations \n"
        "   Center (SOC) at Ext. 9999.\n"
        "2. Do not attempt to investigate or remediate the incident on your own.\n"
        "3. Preserve all evidence (do not delete emails, logs, or files).\n"
        "4. The SOC will triage the incident within 1 hour and escalate as needed.\n"
        "5. Post-incident review will be conducted within 5 business days.\n\n"
        "The company maintains a 24/7 Security Operations Center to monitor threats and respond "
        "to incidents in real-time."
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "5. Remote Work Security", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "Employees working remotely must:\n\n"
        "- Use the company VPN for all work-related activities\n"
        "- Ensure their home Wi-Fi is secured with WPA3 encryption\n"
        "- Lock their screen when stepping away from the computer\n"
        "- Not use public Wi-Fi for accessing company systems\n"
        "- Keep their operating system and software up to date\n"
        "- Use company-approved video conferencing tools only"
    ))
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "6. Compliance and Auditing", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, (
        "The Information Security team conducts quarterly audits of all systems. Employees may be "
        "asked to participate in security assessments or provide access logs. Non-compliance with "
        "these security guidelines will result in a warning for the first offense, mandatory "
        "re-training for the second offense, and disciplinary action for subsequent violations."
    ))

    output_path = SAMPLE_DIR / "security_guidelines.pdf"
    pdf.output(str(output_path))
    print(f"Created: {output_path}")


if __name__ == "__main__":
    print("Creating sample data...\n")
    create_hr_policy_pdf()
    create_product_faq_txt()
    create_onboarding_docx()
    create_security_guidelines_pdf()
    print(f"\nAll sample data created in: {SAMPLE_DIR}")
